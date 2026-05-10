import os
import tempfile
import pytest
from fastapi.testclient import TestClient
from main import app
from docx import Document

client = TestClient(app)

# ============= Basic Functionality Tests =============

def test_root():
    """Test health check endpoint"""
    response = client.get("/")
    assert response.status_code == 200
    assert "Resume2Interview API" in response.json()["message"]


def test_upload_resume_docx():
    """Test successful DOCX resume upload"""
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Software Engineer with 5 years of experience")
    doc.add_paragraph("Skills: Python, FastAPI, PostgreSQL")
    tmp_path = tempfile.mktemp(suffix=".docx")
    doc.save(tmp_path)
    
    try:
        with open(tmp_path, "rb") as f:
            response = client.post(
                "/upload-resume/", 
                files={"file": ("test_resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        assert response.status_code == 200
        data = response.json()
        assert "id" in data
        assert "filename" in data
        assert "parsed" in data
        assert "raw_text" in data["parsed"]
        assert len(data["parsed"]["raw_text"]) > 0
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def test_upload_jd_docx():
    """Test successful DOCX job description upload"""
    doc = Document()
    doc.add_paragraph("Senior QA Engineer")
    doc.add_paragraph("Requirements: 5+ years of testing experience")
    doc.add_paragraph("Skills: Selenium, Python, API Testing")
    tmp_path = tempfile.mktemp(suffix=".docx")
    doc.save(tmp_path)
    
    try:
        with open(tmp_path, "rb") as f:
            response = client.post(
                "/upload-jd/", 
                files={"file": ("test_jd.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        assert response.status_code == 200
        data = response.json()
        assert "id" in data
        assert "parsed" in data
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ============= Error Handling Tests =============

def test_upload_invalid_file_type():
    """Test that invalid file types are rejected with 400"""
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
        tmp.write(b"This is a text file, not a valid resume format")
        tmp.flush()
        tmp_path = tmp.name
    
    try:
        with open(tmp_path, "rb") as f:
            response = client.post(
                "/upload-resume/", 
                files={"file": ("test.txt", f, "text/plain")}
            )
        assert response.status_code == 400
        data = response.json()
        assert "detail" in data
        assert data["detail"]["error_code"] == "INVALID_FILE_TYPE"
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def test_upload_file_too_large():
    """Test that files exceeding size limit are rejected with 413"""
    # Create a file larger than 10 MB
    large_content = b"x" * (11 * 1024 * 1024)  # 11 MB
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(b"%PDF-1.4\n")  # PDF header
        tmp.write(large_content)
        tmp.flush()
        tmp_path = tmp.name
    
    try:
        with open(tmp_path, "rb") as f:
            response = client.post(
                "/upload-resume/", 
                files={"file": ("large.pdf", f, "application/pdf")}
            )
        assert response.status_code == 413
        data = response.json()
        assert "detail" in data
        assert data["detail"]["error_code"] == "FILE_TOO_LARGE"
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def test_upload_empty_file():
    """Test that empty files are rejected"""
    doc = Document()  # Empty document
    tmp_path = tempfile.mktemp(suffix=".docx")
    doc.save(tmp_path)
    
    try:
        with open(tmp_path, "rb") as f:
            response = client.post(
                "/upload-resume/", 
                files={"file": ("empty.docx", f)}
            )
        # Should fail with parsing error or validation error
        assert response.status_code in [422, 400]
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def test_get_resume_not_found():
    """Test 404 for non-existent resume"""
    response = client.get("/resume/999999")
    assert response.status_code == 404
    data = response.json()
    assert "detail" in data
    assert data["detail"]["error_code"] == "NOT_FOUND"


def test_get_jd_not_found():
    """Test 404 for non-existent job description"""
    response = client.get("/jd/999999")
    assert response.status_code == 404
    data = response.json()
    assert "detail" in data
    assert data["detail"]["error_code"] == "NOT_FOUND"


# ============= Integration Tests =============

def test_full_resume_workflow():
    """Test complete workflow: upload → retrieve → verify data"""
    # Upload a resume
    doc = Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("QA Engineer")
    tmp_path = tempfile.mktemp(suffix=".docx")
    doc.save(tmp_path)
    
    try:
        with open(tmp_path, "rb") as f:
            upload_response = client.post(
                "/upload-resume/", 
                files={"file": ("jane_smith.docx", f)}
            )
        assert upload_response.status_code == 200
        resume_id = upload_response.json()["id"]
        
        # Retrieve the resume
        get_response = client.get(f"/resume/{resume_id}")
        assert get_response.status_code == 200
        data = get_response.json()
        assert data["id"] == resume_id
        assert "raw_text" in data
        assert "Jane Smith" in data["raw_text"]
        assert "created_at" in data
        assert "updated_at" in data
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def test_error_response_format():
    """Test that error responses match design specification"""
    response = client.get("/resume/999999")
    assert response.status_code == 404
    data = response.json()
    
    # Verify standardized error format
    assert "detail" in data
    detail = data["detail"]
    assert "error_code" in detail
    assert "message" in detail
    assert "details" in detail or detail.get("details") is None
    assert detail["error_code"] == "NOT_FOUND"


# ============= Validation Tests =============

def test_pdf_magic_number_validation():
    """Test that file content is validated, not just extension"""
    # Create a text file with .pdf extension (should be rejected)
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(b"This is not a real PDF file")
        tmp.flush()
        tmp_path = tmp.name
    
    try:
        with open(tmp_path, "rb") as f:
            response = client.post(
                "/upload-resume/", 
                files={"file": ("fake.pdf", f, "application/pdf")}
            )
        assert response.status_code == 400
        data = response.json()
        assert data["detail"]["error_code"] == "INVALID_FILE_TYPE"
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
